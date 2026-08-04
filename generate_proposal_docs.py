"""
generate_proposal_docs.py — Generate Word (.docx) and LaTeX (.pdf) documents
containing the research proposal steps and H1-H30 hypotheses for academic review.
"""
import os
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

ROOT_DIR = os.path.dirname(os.path.abspath(__file__))

HYPOTHESES = [
    ("H1", "Expiry Day Basis Volatility Elevation", "Basis volatility is identical on Expiry and Control days", "Basis volatility is significantly higher on Expiry days due to VWAP settlement trading", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H2", "Liquidity Impact on Expiry Basis Volatility", "Illiquid and Liquid stocks experience identical basis volatility elevation", "Illiquid stocks exhibit significantly higher basis volatility on Expiry day than Liquid stocks", "Mann-Whitney U Test"),
    ("H3", "Expiry vs Control Spread Widening", "Bid-Ask spreads remain unchanged between Expiry and Control days", "Bid-Ask spreads widen significantly during the 15:00-15:30 window on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H4", "Spread Widening in Illiquid vs Liquid Symbols", "Spread widening on Expiry day is equal across liquidity groups", "Illiquid symbols experience significantly greater percentage spread widening than Liquid symbols", "Mann-Whitney U Test"),
    ("H5", "CLOB Top-5 Depth Erosion on Expiry Day", "Available order book depth at Top-5 price levels is unchanged on Expiry day", "Top-5 limit order book depth erodes significantly on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H6", "Asymmetric Depth Erosion (Illiquid vs Liquid)", "Depth erosion on Expiry day is identical across liquid and illiquid stocks", "Illiquid symbols suffer more severe percentage depth erosion than Liquid symbols", "Mann-Whitney U Test"),
    ("H7", "Final 10-Minute Depth Collapse", "LOB depth remains stable across the entire 30-minute settlement window", "LOB depth drops significantly in the final 10 minutes (15:20-15:30 IST) of Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H8", "Cash Order Flow Imbalance (OFI) Expiry Elevation", "Absolute Cash OFI is identical on Expiry and Control days", "Absolute Cash OFI is significantly elevated on Expiry day due to directional settlement pressure", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H9", "Futures OFI Expiry Elevation", "Absolute Futures OFI is identical on Expiry and Control days", "Absolute Futures OFI is significantly elevated on Expiry day due to contract rolling/closing", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H10", "Cross-Market OFI Synchronization", "Cash and Futures OFI are uncorrelated during the settlement window", "Cash and Futures OFI exhibit strong positive correlation on Expiry day", "Pearson & Spearman Correlation"),
    ("H11", "Kyle's Lambda (Price Impact) Elevation", "Price impact per million INR traded is identical on Expiry and Control days", "Kyle's Lambda (price impact) is significantly higher on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H12", "Illiquid Symbol Price Impact Sensitivity", "Kyle's Lambda elevation is uniform across all Nifty 50 symbols", "Illiquid symbols show a significantly larger increase in Kyle's Lambda on Expiry day", "Mann-Whitney U Test"),
    ("H13", "Proprietary Trader Activity Elevation", "Proprietary trader share of trading volume is unchanged on Expiry day", "Proprietary trader share of volume is significantly higher on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H14", "Custodian Trade Share Stability", "Custodian institutional trading volume share changes dramatically on Expiry day", "Custodian trade share shows minimal deviation between Expiry and Control days", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H15", "Algorithmic Order Contribution Elevation", "Algorithmic order entry share is identical on Expiry and Control days", "Algorithmic order entry share is significantly higher during Expiry settlement", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H16", "Algo IOC Aggressiveness Elevation", "Immediate-Or-Cancel (IOC) order usage by algos is identical across days", "Algorithmic IOC usage increases significantly on Expiry day for urgent execution", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H17", "Expiry Day Order Cancellation Ratio Elevation", "Order cancellation-to-execution ratio is identical on Expiry and Control days", "Order cancellation ratio is significantly higher on Expiry day due to fleeting liquidity", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H18", "High-Frequency Quote Modification Intensity", "Order modification frequency per minute is unchanged on Expiry day", "Quote modification frequency is significantly elevated on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H19", "Iceberg Order Hidden Volume Contribution", "Hidden iceberg order volume is identical on Expiry and Control days", "Hidden iceberg order volume is significantly higher on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H20", "Bloomberg Roll Direction Predicts Cash VWAP Drift", "Bloomberg roll pressure direction has no predictive power over terminal VWAP drift (<50% accuracy)", "Bloomberg roll direction correctly predicts the sign of cash VWAP drift (>50% accuracy)", "Binomial Test (H0: p=0.5)"),
    ("H21", "Roll Intensity Correlation with Magnitude of Cash VWAP Drift", "No correlation exists between roll intensity and magnitude of cash VWAP drift", "Roll intensity is positively correlated with absolute cash VWAP terminal drift", "Spearman Rank Correlation"),
    ("H22", "Roll Direction Predicts Final CLOB Book Asymmetry", "Bloomberg roll direction does not predict final limit order book imbalance (<50% accuracy)", "Bloomberg roll direction predicts whether final CLOB imbalance is bid-heavy or ask-heavy", "Binomial Test (H0: p=0.5)"),
    ("H23", "Cost of Carry Mispricing Correlation with Cash OFI", "No correlation exists between basis mispricing and Cash Order Flow Imbalance", "Absolute cost-of-carry mispricing correlates positively with Cash OFI", "Spearman Rank Correlation"),
    ("H24", "Cash VWAP Terminal Acceleration", "VWAP drift is linear across the 30-minute settlement window", "VWAP drift accelerates significantly in the final 10 minutes (15:20-15:30 IST)", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H25", "Futures Volume Concentration in Final 15 Minutes", "Futures trading volume is evenly distributed across the 15:00-15:30 window", "Futures volume concentrates significantly in the 15:15-15:30 IST window", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H26", "Expiry Day Turnover Elevation", "Combined Cash and Futures INR turnover is identical on Expiry and Control days", "Combined INR turnover is significantly elevated on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H27", "Cross-Market Spread Convergence at 15:30 IST", "Futures-Cash basis spread at 15:30 IST is identical on Expiry and Control days", "Futures-Cash basis converges closer to zero at 15:30 IST on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H28", "Algorithmic Share of Trade Volume Elevation", "Algo share of executed trade volume is unchanged on Expiry day", "Algorithmic execution volume share increases significantly on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H29", "Order Book Imbalance Persistence", "Intraday autocorrelation of OFI is identical on Expiry and Control days", "OFI autocorrelation is significantly stronger on Expiry day due to directional push", "Paired t-test & Wilcoxon Signed-Rank"),
    ("H30", "Expiry Day Intraday Price Reversal Post-15:15 IST", "Price movement post-15:15 is independent of 15:00-15:15 movement", "Prices exhibit significant mean-reverting reversal post-15:15 IST on Expiry day", "Paired t-test & Wilcoxon Signed-Rank"),
]

STEPS_TEXT = [
    (
        "Step 1: Sample Selection & Experimental Design (10 Symbols x 24 Dates)",
        [
            "Target Asset Universe: Select 5 Liquid Nifty 50 stocks (RELIANCE, TCS, INFY, HDFCBANK, ICICIBANK) and 5 Illiquid Nifty 50 stocks (GRASIM, HEROMOTOCO, SHREEVECEM, UPL, TECHM).",
            "Experimental Pairing: Match 12 Monthly Expiry Thursdays in 2022 against 12 Control Thursdays (typically 2 weeks prior to expiry) to isolate expiry-day effects from day-of-week seasonality.",
            "Settlement Focus Window: Extract tick-by-tick cash market and derivatives market data during the NSE VWAP settlement window (15:00 to 15:30 IST)."
        ]
    ),
    (
        "Step 2: High-Speed Binary Tick Data Parsing (Stage 1)",
        [
            "Input Data Format: Ingest fixed-width binary text files (.DAT.gz) from NSE for CASH_Orders, CASH_Trades, FAO_Orders, and FAO_Trades.",
            "Predicate Pushdown: Filter raw strings on symbol substrings and instrument types ('EQ' for Cash, 'FUTSTK' for Futures) before extracting schema columns.",
            "Column Extraction: Extract 17 standardized fields including Order/Trade numbers, timestamps (Jiffies), prices, volumes, client identity, and algo indicator flags, saving as partitioned Parquet."
        ]
    ),
    (
        "Step 3: Temporal & Microstructure Enrichment (Stage 2)",
        [
            "Timestamp Normalization: Convert NSE Jiffies (1/65536th of a second since epoch) into standard Python/Spark timestamps and HH:mm:ss time buckets.",
            "Currency & Flag Mapping: Convert Paise to INR Rupees (/ 100.0), flag the 30-minute settlement window (15:00-15:30 IST), and identify Expiry vs. Control dates.",
            "Participant & Algo Tagging: Categorize Client Identity into Custodian (Institutional), Proprietary, and NCNP (Retail), and classify Algo Indicator flags into Algo vs. Non-Algo execution.",
            "Dynamic Partitioning: Save enriched datasets partitioned by ['symbol', 'trade_date'] to enable instant single-symbol filter pushdown during downstream CLOB replay."
        ]
    ),
    (
        "Step 4: Trade-Level Microstructure Metrics Computation (Stage 3)",
        [
            "A1 (VWAP Trajectory): Compute cumulative and rolling 1-minute Volume-Weighted Average Price (VWAP) trajectories across 15:00-15:30 IST.",
            "A2 (Basis Divergence): Compute intraday Futures-Cash basis spread and Basis Volatility (standard deviation of basis).",
            "A3-A7 (Participant & Algo Profiling): Quantify Prop/Custodian trade shares, Algo volume contribution, Immediate-Or-Cancel (IOC) aggressiveness, cancellation-to-trade ratios, and iceberg hidden order volume."
        ]
    ),
    (
        "Step 5: High-Performance Limit Order Book (CLOB) Reconstruction (Stage 4)",
        [
            "Chronological Event Merging: Merge order entry/modify/cancel events with trade execution events into a unified chronological stream.",
            "Data-Oriented Design (C++): Replay events through a cache-line aligned C++ order book engine (FlatPriceBook with O(log N) price level lookup and zero heap fragmentation).",
            "1-Second Depth Snapshots: Emit Top-10 bid/ask price levels and available shares at every second from 15:00:00 to 15:30:00 IST (1,800 snapshots per symbol/date)."
        ]
    ),
    (
        "Step 6: Limit Order Book Liquidity & Price Impact Analysis (Stage 5)",
        [
            "B1-B2 (Spread & Depth Dynamics): Calculate percentage bid-ask spread widening and Top-5 / Top-10 order book depth erosion across the settlement window.",
            "B3-B5 (OFI, Kyle's Lambda & Asymmetry): Calculate Order Flow Imbalance (OFI), estimate Kyle's Lambda (price impact per million INR traded via OLS regression), and measure Bid-vs-Ask book asymmetry."
        ]
    ),
    (
        "Step 7: Bloomberg Terminal Integration & Cross-Market Validation (Stage 6)",
        [
            "C1 (Roll Pressure): Ingest Bloomberg calendar spread data to quantify institutional rollover urgency and predict net directional punch ('UP' or 'DOWN').",
            "C2-C3 (Cost-of-Carry & Cross-Validation): Calculate implied repo rates and test whether Bloomberg roll direction predicts actual NSE Cash VWAP terminal drift and CLOB book imbalance.",
            "C4 (Event Study): Evaluate price trajectories around block trade executions during the final 15 minutes."
        ]
    ),
    (
        "Step 8: Formal Hypothesis Testing Engine & Statistical Reporting (Stage 7)",
        [
            "Statistical Test Execution: Evaluate all 30 formal hypotheses (H1-H30) using Paired t-tests, Wilcoxon Signed-Rank tests, Mann-Whitney U tests, OLS regression, Spearman rank correlation, and Binomial tests.",
            "FDR Multiple-Testing Correction: Apply Benjamini-Hochberg False Discovery Rate correction across all p-values to control Type I errors.",
            "Artifact Compilation: Automatically generate publication-quality figures (fig1 to fig10), summary statistical tables, and the final Markdown research paper."
        ]
    ),
    (
        "Step 9: Modular Staged Execution & Scale-Up Strategy",
        [
            "Staged Workflow: Start processing with 1 or 2 symbols/dates for rapid initial verification of the entire pipeline.",
            "Incremental Scaling: Copy additional raw data files into 'data/raw/' at any time; pipeline automatically skips previously parsed/built snapshots and dynamically updates partitions without wiping existing data."
        ]
    )
]

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_word_document():
    doc_path = os.path.join(ROOT_DIR, "NSE_Expiry_Day_Analysis_Proposal_and_Steps.docx")
    print(f"[DOCX] Creating Word Document: {doc_path} ...")
    doc = docx.Document()

    # Page Margins (1 inch)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NSE Expiry Day Dynamics & VWAP Settlement Anomalies")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 54, 93) # Deep Navy

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Research Methodology Workflow, Processing Steps & Complete Hypotheses (H1–H30)")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(74, 85, 104)

    doc.add_paragraph() # spacing

    # Section 1: Rough Processing Steps
    h1 = doc.add_heading("1. Methodology & Rough Processing Steps (Pipeline Overview)", level=1)
    h1.style.font.color.rgb = RGBColor(26, 54, 93)

    p_intro = doc.add_paragraph(
        "This research examines the microstructure of 10 Nifty 50 stocks (5 liquid, 5 illiquid) and their corresponding "
        "futures contracts on the National Stock Exchange (NSE) during the final 30-minute settlement window (15:00–15:30 IST) "
        "across 12 Monthly Expiry Thursdays and 12 matched Control trading days in 2022. Below is the step-by-step processing workflow:"
    )
    p_intro.style.font.name = "Calibri"
    p_intro.style.font.size = Pt(11)

    for title, bullets in STEPS_TEXT:
        p_step = doc.add_heading(title, level=2)
        p_step.style.font.color.rgb = RGBColor(43, 108, 176)
        for b in bullets:
            p_b = doc.add_paragraph(b, style="List Bullet")
            p_b.style.font.name = "Calibri"
            p_b.style.font.size = Pt(10.5)

    doc.add_paragraph() # spacing

    # Section 2: 30 Formal Hypotheses
    h2 = doc.add_heading("2. Complete List of Tested Hypotheses (H1 – H30)", level=1)
    h2.style.font.color.rgb = RGBColor(26, 54, 93)

    p_h_intro = doc.add_paragraph(
        "The table below outlines all 30 formal statistical hypotheses evaluated in Stage 7 of the pipeline, "
        "along with their null hypothesis (H0), alternative hypothesis (H1), and statistical testing methodology."
    )
    p_h_intro.style.font.name = "Calibri"
    p_h_intro.style.font.size = Pt(11)

    # Table creation
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(0.6), Inches(1.5), Inches(1.7), Inches(1.7), Inches(1.0)]
    headers = ["ID", "Hypothesis Name", "Null Hypothesis (H0)", "Alternative Hypothesis (H1)", "Statistical Test"]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i], 120, 120, 150, 150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (hid, hname, h0, h1_text, test_name) in enumerate(HYPOTHESES):
        row_cells = table.add_row().cells
        bg_color = "F7FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate([hid, hname, h0, h1_text, test_name]):
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], 100, 100, 120, 120)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row_cells[i].paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            if i == 0:
                run.font.bold = True

    doc.save(doc_path)
    print(f"[DONE] Saved Word document to: {doc_path}")

def create_latex_document():
    tex_path = os.path.join(ROOT_DIR, "NSE_Expiry_Day_Analysis_Proposal_and_Steps.tex")
    pdf_path = os.path.join(ROOT_DIR, "NSE_Expiry_Day_Analysis_Proposal_and_Steps.pdf")
    print(f"[LATEX] Creating LaTeX Document: {tex_path} ...")

    lines = [
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage[top=1in,bottom=1in,left=0.8in,right=0.8in]{geometry}",
        r"\usepackage{booktabs}",
        r"\usepackage{longtable}",
        r"\usepackage{array}",
        r"\usepackage{enumitem}",
        r"\usepackage{xcolor}",
        r"\usepackage{titlesec}",
        r"\usepackage{hyperref}",
        r"\definecolor{deepnavy}{RGB}{26,54,93}",
        r"\definecolor{sectionblue}{RGB}{43,108,176}",
        r"\titleformat{\section}{\Large\bfseries\color{deepnavy}}{\thesection}{1em}{}",
        r"\titleformat{\subsection}{\large\bfseries\color{sectionblue}}{\thesubsection}{1em}{}",
        r"\hypersetup{colorlinks=true,linkcolor=deepnavy,urlcolor=sectionblue}",
        r"\begin{document}",
        r"\begin{center}",
        r"{\LARGE \textbf{\color{deepnavy} NSE Expiry Day Dynamics \& VWAP Settlement Anomalies}\\[0.4cm]}",
        r"{\large \textit{Research Methodology Workflow, Processing Steps \& Complete Hypotheses (H1--H30)}\\[0.5cm]}",
        r"\end{center}",
        r"\vspace{0.3cm}",
        r"\section{Methodology \& Rough Processing Steps (Pipeline Overview)}",
        r"This research examines the microstructure of 10 Nifty 50 stocks (5 liquid, 5 illiquid) and their corresponding "
        r"futures contracts on the National Stock Exchange (NSE) during the final 30-minute settlement window (15:00--15:30 IST) "
        r"across 12 Monthly Expiry Thursdays and 12 matched Control trading days in 2022. Below is the step-by-step processing workflow:",
        r"\vspace{0.2cm}"
    ]

    for title, bullets in STEPS_TEXT:
        clean_title = title.replace("&", r"\&")
        lines.append(rf"\subsection*{{{clean_title}}}")
        lines.append(r"\begin{itemize}[noitemsep,topsep=2pt]")
        for b in bullets:
            clean_b = b.replace("&", r"\&").replace("%", r"\%").replace("_", r"\_")
            lines.append(rf"  \item {clean_b}")
        lines.append(r"\end{itemize}")
        lines.append(r"\vspace{0.15cm}")

    lines.append(r"\newpage")
    lines.append(r"\section{Complete List of Tested Hypotheses (H1 -- H30)}")
    lines.append(
        r"The table below outlines all 30 formal statistical hypotheses evaluated in Stage 7 of the pipeline, "
        r"along with their null hypothesis ($H_0$), alternative hypothesis ($H_1$), and statistical testing methodology."
    )
    lines.append(r"\vspace{0.3cm}")

    lines.append(r"\begin{longtable}{p{0.8cm} p{3.2cm} p{4.2cm} p{4.2cm} p{2.8cm}}")
    lines.append(r"\toprule")
    lines.append(r"\textbf{ID} & \textbf{Hypothesis Name} & \textbf{Null Hypothesis ($H_0$)} & \textbf{Alternative Hypothesis ($H_1$)} & \textbf{Statistical Test} \\")
    lines.append(r"\midrule")
    lines.append(r"\endfirsthead")
    lines.append(r"\toprule")
    lines.append(r"\textbf{ID} & \textbf{Hypothesis Name} & \textbf{Null Hypothesis ($H_0$)} & \textbf{Alternative Hypothesis ($H_1$)} & \textbf{Statistical Test} \\")
    lines.append(r"\midrule")
    lines.append(r"\endhead")
    lines.append(r"\bottomrule")
    lines.append(r"\endfoot")
    lines.append(r"\bottomrule")
    lines.append(r"\endlastfoot")

    for hid, hname, h0, h1_text, test_name in HYPOTHESES:
        clean_name = hname.replace("&", r"\&").replace("_", r"\_")
        clean_h0 = h0.replace("&", r"\&").replace("%", r"\%").replace("_", r"\_").replace("<", r"$<$").replace(">", r"$>$")
        clean_h1 = h1_text.replace("&", r"\&").replace("%", r"\%").replace("_", r"\_").replace("<", r"$<$").replace(">", r"$>$")
        clean_test = test_name.replace("&", r"\&").replace("_", r"\_")
        lines.append(rf"\textbf{{{hid}}} & {clean_name} & {clean_h0} & {clean_h1} & {clean_test} \\")
        lines.append(r"\addlinespace[3pt]")

    lines.append(r"\end{longtable}")
    lines.append(r"\end{document}")

    with open(tex_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[DONE] Saved LaTeX source to: {tex_path}")

    # Compile LaTeX to PDF
    pdflatex_exe = r"C:\Users\arnav\AppData\Roaming\TinyTeX\bin\windows\pdflatex.exe"
    if not os.path.exists(pdflatex_exe):
        pdflatex_exe = "pdflatex"

    print(f"[COMPILE] Compiling LaTeX to PDF via {pdflatex_exe} ...")
    try:
        # Run pdflatex twice for longtable headers/page numbering
        for _ in range(2):
            subprocess.run(
                [pdflatex_exe, "-interaction=nonstopmode", "NSE_Expiry_Day_Analysis_Proposal_and_Steps.tex"],
                cwd=ROOT_DIR,
                check=True,
                stdout=subprocess.DEVNULL
            )
        print(f"[SUCCESS] Compiled PDF successfully: {pdf_path}")
    except Exception as e:
        print(f"[WARN] Error compiling LaTeX: {e}")

if __name__ == "__main__":
    create_word_document()
    create_latex_document()
